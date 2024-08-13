# Import necessary libraries
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split, cross_val_score
from sklearn.metrics import r2_score, mean_squared_error, mean_absolute_error
from sklearn.ensemble import GradientBoostingRegressor
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Load the dataset
df = pd.read_csv('SolarPrediction.csv')
print(df.head())

# Initialize a Word Document
doc = Document()
doc.add_heading('Solar Radiation Prediction Results - Gradient Boosting Regressor', 0)

# Exploratory Data Analysis (EDA)
plt.figure(figsize=(12, 6))
sns.histplot(df['UNIXTime'], kde=True)
plt.title('Distribution of UNIXTime')
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('Distribution of UNIXTime', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Drop unnecessary columns
df = df.drop(['Data', 'Time', 'TimeSunRise', 'TimeSunSet'], axis=1)
print(df.head())

# Heatmap of correlations
plt.figure(figsize=(10, 10))
sns.heatmap(df.corr(), annot=True, cmap='coolwarm', center=0)
plt.title('Correlation Heatmap')
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('Correlation Heatmap', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Split data into features and target variable
X = df.drop('Radiation', axis=1)
y = df['Radiation']

# Split data into training and testing sets
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.30, random_state=101)

# Define the Gradient Boosting Regressor model
gb_model = GradientBoostingRegressor(n_estimators=100, learning_rate=0.1, max_depth=3, random_state=101)

# Train the model
gb_model.fit(X_train, y_train)

# Make predictions
y_pred = gb_model.predict(X_test)

# Evaluate the model
r2 = r2_score(y_test, y_pred)
mse = mean_squared_error(y_test, y_pred)
rmse = np.sqrt(mse)
mae = mean_absolute_error(y_test, y_pred)

doc.add_heading('Regression Metrics', level=1)
doc.add_paragraph(f"R2 Score: {r2}")
doc.add_paragraph(f"Mean Squared Error (MSE): {mse}")
doc.add_paragraph(f"Root Mean Squared Error (RMSE): {rmse}")
doc.add_paragraph(f"Mean Absolute Error (MAE): {mae}")

# Cross-validation
cv_results = cross_val_score(gb_model, X, y, cv=10, scoring='neg_mean_squared_error')
doc.add_heading('Cross-Validation Results', level=1)
doc.add_paragraph(f"Cross-Validation MSE Scores: {-cv_results}")
doc.add_paragraph(f"Mean CV MSE Score: {-cv_results.mean()}")
doc.add_paragraph(f"Standard Deviation of CV MSE Scores: {cv_results.std()}")

# Plot true vs predicted values
plt.figure(figsize=(10, 6))
plt.scatter(y_test, y_pred, alpha=0.5)
plt.plot([y.min(), y.max()], [y.min(), y.max()], 'r--')
plt.xlabel('True Values')
plt.ylabel('Predicted Values')
plt.title('True vs Predicted Radiation Values')
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('True vs Predicted Radiation Values', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Plot residuals
residuals = y_test - y_pred
plt.figure(figsize=(10, 6))
sns.histplot(residuals, kde=True)
plt.title('Distribution of Residuals')
plt.xlabel('Residual')
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('Distribution of Residuals', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Save the document
doc.save('Solar_Radiation_Prediction_GradientBoosting_Regressor_Results.docx')