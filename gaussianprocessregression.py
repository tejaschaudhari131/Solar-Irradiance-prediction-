# Import necessary libraries
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split, cross_val_score
from sklearn.metrics import mean_squared_error, r2_score
from sklearn.gaussian_process import GaussianProcessRegressor
from sklearn.gaussian_process.kernels import RBF, ConstantKernel as C
from sklearn.preprocessing import StandardScaler
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Load the dataset
df = pd.read_csv('SolarPrediction.csv')
print(df.head())

# Initialize a Word Document
doc = Document()
doc.add_heading('Solar Radiation Prediction Results - Gaussian Process Regression', 0)

# Exploratory Data Analysis (EDA)
plt.figure(figsize=(12, 6))
sns.histplot(df['UNIXTime'], kde=True)
plt.title('Distribution of UNIXTime')
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('Distribution of UNIXTime', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Drop unnecessary columns
df = df.drop(['Data', 'Time', 'TimeSunRise', 'TimeSunSet'], axis=1)
print(df.head())

# Heatmap of correlations
plt.figure(figsize=(10, 10))
sns.heatmap(df.corr(), annot=True, cmap='coolwarm', center=0)
plt.title('Correlation Heatmap')
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('Correlation Heatmap', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Split data into features and target variable
X = df.drop('Radiation', axis=1)
y = df['Radiation']

# Split data into training and testing sets
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.30, random_state=101)

# Scale the features
scaler = StandardScaler()
X_train_scaled = scaler.fit_transform(X_train)
X_test_scaled = scaler.transform(X_test)

# Define the Gaussian Process Regression model
kernel = C(1.0, (1e-3, 1e3)) * RBF(10, (1e-2, 1e2))
gpr = GaussianProcessRegressor(kernel=kernel, n_restarts_optimizer=10, random_state=101)

# Train the model
gpr.fit(X_train_scaled, y_train)

# Make predictions
y_pred, sigma = gpr.predict(X_test_scaled, return_std=True)

# Evaluate the model
mse = mean_squared_error(y_test, y_pred)
rmse = np.sqrt(mse)
r2 = r2_score(y_test, y_pred)

doc.add_heading('Regression Metrics', level=1)
doc.add_paragraph(f"Mean Squared Error: {mse}")
doc.add_paragraph(f"Root Mean Squared Error: {rmse}")
doc.add_paragraph(f"R2 Score: {r2}")

# Cross-validation
cv_results = cross_val_score(gpr, X_train_scaled, y_train, cv=5, scoring='neg_mean_squared_error')
cv_results = -cv_results  # Convert to positive MSE
doc.add_heading('Cross-Validation Results', level=1)
doc.add_paragraph(f"Cross-Validation MSE Scores: {cv_results}")
doc.add_paragraph(f"Mean CV MSE Score: {cv_results.mean()}")
doc.add_paragraph(f"Standard Deviation of CV MSE Scores: {cv_results.std()}")

# Plot true vs predicted values
plt.figure(figsize=(10, 6))
plt.scatter(y_test, y_pred, alpha=0.5)
plt.plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()], 'r--', lw=2)
plt.xlabel('True Values')
plt.ylabel('Predictions')
plt.title('True vs Predicted Radiation Values')
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('True vs Predicted Radiation Values', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Plot residuals
residuals = y_test - y_pred
plt.figure(figsize=(10, 6))
sns.histplot(residuals, kde=True)
plt.title('Distribution of Residuals')
plt.xlabel('Residual')
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('Distribution of Residuals', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Plot predictions with uncertainty
plt.figure(figsize=(10, 6))
plt.scatter(y_test, y_pred, alpha=0.5)
plt.errorbar(y_test, y_pred, yerr=sigma, fmt='o', alpha=0.2)
plt.plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()], 'r--', lw=2)
plt.xlabel('True Values')
plt.ylabel('Predictions')
plt.title('True vs Predicted Radiation Values with Uncertainty')
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('True vs Predicted Radiation Values with Uncertainty', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Save the document5erujnl.hm,
doc.save('GaussianProcessRegression_Results.docx')