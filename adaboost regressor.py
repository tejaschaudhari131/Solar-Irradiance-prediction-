# Import necessary libraries
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split, cross_val_score
from sklearn.metrics import r2_score, mean_squared_error, mean_absolute_error
from sklearn.ensemble import AdaBoostRegressor
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler
from sklearn.compose import ColumnTransformer
from sklearn.impute import SimpleImputer
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Load the dataset
df = pd.read_csv('SolarPrediction.csv')
print(df.head())

# Initialize a Word Document
doc = Document()
doc.add_heading('Solar Radiation Prediction Results', 0)

# Drop unnecessary columns
df = df.drop(['Data', 'Time', 'TimeSunRise', 'TimeSunSet'], axis=1)

# Feature Engineering
def create_time_features(df):
    df['datetime'] = pd.to_datetime(df['UNIXTime'], unit='s')
    df['hour'] = df['datetime'].dt.hour
    df['dayofyear'] = df['datetime'].dt.dayofyear
    df['dayofweek'] = df['datetime'].dt.dayofweek
    df['month'] = df['datetime'].dt.month
    df = df.drop(['UNIXTime', 'datetime'], axis=1)
    return df

df = create_time_features(df)
print(df.head())

# Heatmap of correlations
plt.figure(figsize=(10, 10))
sns.heatmap(df.corr(), annot=True, cmap='coolwarm', center=0)
plt.title('Correlation Heatmap')
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('Correlation Heatmap', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Split data into features and target variable
X = df.drop('Radiation', axis=1)
y = df['Radiation']

# Split data into training and testing sets
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.30, random_state=101)

# Preprocessing Pipeline
numeric_features = X_train.columns.tolist()
numeric_transformer = Pipeline(steps=[
    ('imputer', SimpleImputer(strategy='median')),
    ('scaler', StandardScaler())
])

preprocessor = ColumnTransformer(
    transformers=[
        ('num', numeric_transformer, numeric_features)
    ])

# Define the AdaBoost model
adaboost_model = AdaBoostRegressor(random_state=101)

# Create the final pipeline with AdaBoost only
model_pipeline = Pipeline(steps=[
    ('preprocessor', preprocessor),
    ('adaboost', adaboost_model)
])

# Fit the model
model_pipeline.fit(X_train, y_train)

# Evaluate the model
y_pred = model_pipeline.predict(X_test)
r2 = r2_score(y_test, y_pred)
mse = mean_squared_error(y_test, y_pred)
rmse = np.sqrt(mse)
mae = mean_absolute_error(y_test, y_pred)

doc.add_heading('Regression Metrics', level=1)
doc.add_paragraph(f"R2 Score: {r2}")
doc.add_paragraph(f"Mean Squared Error (MSE): {mse}")
doc.add_paragraph(f"Root Mean Squared Error (RMSE): {rmse}")
doc.add_paragraph(f"Mean Absolute Error (MAE): {mae}")

# Cross-validation
cv_results = cross_val_score(model_pipeline, X, y, cv=10, scoring='r2')
doc.add_heading('Cross-Validation Results', level=1)
doc.add_paragraph(f"Cross-Validation R2 Scores: {cv_results}")
doc.add_paragraph(f"Mean CV R2 Score: {cv_results.mean()}")
doc.add_paragraph(f"Standard Deviation of CV R2 Scores: {cv_results.std()}")

# Plot true vs predicted values
plt.figure(figsize=(10, 6))
plt.scatter(y_test, y_pred, alpha=0.5)
plt.plot([y.min(), y.max()], [y.min(), y.max()], 'r--')
plt.xlabel('True Values')
plt.ylabel('Predicted Values')
plt.title('True vs Predicted Radiation Values')
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('True vs Predicted Radiation Values', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Plot residuals
residuals = y_test - y_pred
plt.figure(figsize=(10, 6))
sns.histplot(residuals, kde=True)
plt.title('Distribution of Residuals')
plt.xlabel('Residual')
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('Distribution of Residuals', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Save the document
doc.save('Solar_Radiation_Prediction_ADABOOST_Results.docx')
