# Import necessary libraries
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split, cross_val_score
from sklearn.metrics import r2_score, mean_squared_error, mean_absolute_error, accuracy_score, f1_score, precision_score, recall_score
import xgboost as xgb
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Load the dataset
df = pd.read_excel('SolarPrediction data.xlsx')
print(df.head())

# Initialize a Word Document
doc = Document()
doc.add_heading('Solar Radiation Prediction Results', 0)

# Exploratory Data Analysis (EDA)
plt.figure(figsize=(12, 6))
sns.histplot(df['UNIXTime'], kde=True)
plt.title('Distribution of UNIXTime')
# Save the plot to a BytesIO object
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('Distribution of UNIXTime', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Drop unnecessary columns
df = df.drop(['Data', 'Time', 'TimeSunRise', 'TimeSunSet'], axis=1)
print(df.head())

# Heatmap of correlations
plt.figure(figsize=(10, 10))
sns.heatmap(df.corr(), annot=True, cmap='coolwarm', center=0)
plt.title('Correlation Heatmap')
# Save the plot to a BytesIO object
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('Correlation Heatmap', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Pairplot of the dataframe
sns.pairplot(df)
# Save the plot to a BytesIO object
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('Pairplot of DataFrame', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Convert Radiation to a binary classification problem
df['Radiation_Class'] = np.where(df['Radiation'] > df['Radiation'].median(), 1, 0)

# Split data into features and target variable
X = df.drop(['Radiation', 'Radiation_Class'], axis=1)
y = df['Radiation_Class']

# Split data into training and testing sets
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.30, random_state=101)

# Convert the data into DMatrix format for xgboost
dtrain = xgb.DMatrix(X_train, label=y_train)
dtest = xgb.DMatrix(X_test, label=y_test)

# Set parameters for xgboost with CPU support
params = {
    'objective': 'binary:logistic',
    'tree_method': 'hist',  # Use 'hist' instead of 'gpu_hist'
    'random_state': 101
}

# Train the model
xg_reg = xgb.train(params, dtrain, num_boost_round=50)

# Evaluate the model
y_pred = xg_reg.predict(dtest)
y_pred_class = (y_pred > 0.5).astype(int)

accuracy = accuracy_score(y_test, y_pred_class)
precision = precision_score(y_test, y_pred_class)
recall = recall_score(y_test, y_pred_class)
f1 = f1_score(y_test, y_pred_class)

doc.add_heading('Classification Metrics', level=1)
doc.add_paragraph(f"Accuracy: {accuracy}")
doc.add_paragraph(f"Precision: {precision}")
doc.add_paragraph(f"Recall: {recall}")
doc.add_paragraph(f"F1 Score: {f1}")

# Cross-validation
cv_results = xgb.cv(params, xgb.DMatrix(X, label=y), nfold=10, num_boost_round=50, metrics='rmse', as_pandas=True, seed=101)
doc.add_heading('Cross-Validation Results', level=1)
doc.add_paragraph(f"Cross-Validation RMSE Scores: {cv_results['test-rmse-mean']}")
doc.add_paragraph(f"Mean CV RMSE Score: {cv_results['test-rmse-mean'].mean()}")
doc.add_paragraph(f"Standard Deviation of CV RMSE Scores: {cv_results['test-rmse-std'].mean()}")

# Plot true vs predicted values
plt.figure(figsize=(10, 6))
plt.scatter(y_test, y_pred, alpha=0.5)
plt.plot([y.min(), y.max()], [y.min(), y.max()], 'r--')
plt.xlabel('True Values')
plt.ylabel('Predicted Values')
plt.title('True vs Predicted Radiation Values')
# Save the plot to a BytesIO object
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('True vs Predicted Radiation Values', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Plot residuals
residuals = y_test - y_pred
plt.figure(figsize=(10, 6))
sns.histplot(residuals, kde=True)
plt.title('Distribution of Residuals')
plt.xlabel('Residual')
# Save the plot to a BytesIO object
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('Distribution of Residuals', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Save the document
doc.save('Solar_Radiation_Prediction_Results_XGBOOST.docx')
