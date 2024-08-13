# Import necessary libraries
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split, cross_val_score
from sklearn.metrics import accuracy_score, f1_score, precision_score, recall_score
from sklearn.ensemble import GradientBoostingClassifier
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Load the dataset
df = pd.read_csv('SolarPrediction.csv')
print(df.head())

# Initialize a Word Document
doc = Document()
doc.add_heading('Solar Radiation Prediction Results - Gradient Boosting Classifier', 0)

# Exploratory Data Analysis (EDA)
plt.figure(figsize=(12, 6))
sns.histplot(df['UNIXTime'], kde=True)
plt.title('Distribution of UNIXTime')
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('Distribution of UNIXTime', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Drop unnecessary columns
df = df.drop(['Data', 'Time', 'TimeSunRise', 'TimeSunSet'], axis=1)
print(df.head())

# Heatmap of correlations
plt.figure(figsize=(10, 10))
sns.heatmap(df.corr(), annot=True, cmap='coolwarm', center=0)
plt.title('Correlation Heatmap')
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('Correlation Heatmap', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Convert Radiation to a binary classification problem
df['Radiation_Class'] = np.where(df['Radiation'] > df['Radiation'].median(), 1, 0)

# Split data into features and target variable
X = df.drop(['Radiation', 'Radiation_Class'], axis=1)
y = df['Radiation_Class']

# Split data into training and testing sets
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.30, random_state=101)

# Define the Gradient Boosting Classifier model
gb_model = GradientBoostingClassifier(n_estimators=100, learning_rate=0.1, max_depth=3, random_state=101)

# Train the model
gb_model.fit(X_train, y_train)

# Make predictions
y_pred = gb_model.predict(X_test)

# Evaluate the model
accuracy = accuracy_score(y_test, y_pred)
precision = precision_score(y_test, y_pred)
recall = recall_score(y_test, y_pred)
f1 = f1_score(y_test, y_pred)

doc.add_heading('Classification Metrics', level=1)
doc.add_paragraph(f"Accuracy: {accuracy}")
doc.add_paragraph(f"Precision: {precision}")
doc.add_paragraph(f"Recall: {recall}")
doc.add_paragraph(f"F1 Score: {f1}")

# Cross-validation
cv_results = cross_val_score(gb_model, X, y, cv=10, scoring='accuracy')
doc.add_heading('Cross-Validation Results', level=1)
doc.add_paragraph(f"Cross-Validation Accuracy Scores: {cv_results}")
doc.add_paragraph(f"Mean CV Accuracy Score: {cv_results.mean()}")
doc.add_paragraph(f"Standard Deviation of CV Accuracy Scores: {cv_results.std()}")

# Plot true vs predicted values
plt.figure(figsize=(10, 6))
plt.scatter(range(len(y_test)), y_test, alpha=0.5, label='True')
plt.scatter(range(len(y_pred)), y_pred, alpha=0.5, label='Predicted')
plt.legend()
plt.xlabel('Sample Index')
plt.ylabel('Class Label')
plt.title('True vs Predicted Radiation Classes')
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('True vs Predicted Radiation Classes', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Plot residuals
residuals = y_test - y_pred
plt.figure(figsize=(10, 6))
sns.histplot(residuals, kde=True)
plt.title('Distribution of Residuals (True - Predicted)')
plt.xlabel('Residual')
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
plt.close()
img_stream.seek(0)
doc.add_heading('Distribution of Residuals', level=1)
doc.add_picture(img_stream, width=Inches(5))

# Save the document
doc.save('GradientBoosting_Classifier_Results.docx')